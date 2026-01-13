from docx import Document

docx_path = r'c:\ProjectCode\stockmanager\docs\vc투자메모.docx'
doc = Document(docx_path)

print(f"Total paragraphs: {len(doc.paragraphs)}\n")
print("="*80)

for i, para in enumerate(doc.paragraphs[:100], 1):
    text = para.text.strip()
    if text:
        style = para.style.name
        print(f"{i}. [{style}] {text}")

print("\n" + "="*80)
print("\nTables found:", len(doc.tables))
if doc.tables:
    print("\nFirst table preview:")
    table = doc.tables[0]
    for row in table.rows[:5]:
        print(" | ".join([cell.text.strip() for cell in row.cells]))

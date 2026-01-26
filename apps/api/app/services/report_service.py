import os
import datetime
import time
import re
import google.generativeai as genai
from sqlalchemy.orm import Session
from sqlalchemy import text
from typing import List
import xml.etree.ElementTree as ET
import urllib.parse
import requests
import unicodedata
import concurrent.futures
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# Import settings to get API key
from ..config import settings

def _append_report_log(report_id: int, message: str) -> None:
    try:
        root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../.."))
        log_path = os.path.join(root_dir, "artifacts", "report_debug.log")
        ts = datetime.datetime.now().isoformat()
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"{ts} | report_id={report_id} | {message}\n")
    except Exception:
        pass

def fetch_google_news(query: str, limit: int | None = None) -> str:
    """Fetch recent news from Google News RSS"""
    try:
        encoded_query = urllib.parse.quote(query)
        url = f"https://news.google.com/rss/search?q={encoded_query}&hl=ko&gl=KR&ceid=KR:ko"
        response = requests.get(url, timeout=5)
        
        if response.status_code != 200:
            return "뉴스 데이터를 불러오지 못했습니다."
            
        root = ET.fromstring(response.content)
        items = root.findall('.//item')
        
        news_list = []
        if limit is None:
            target_items = items
        else:
            target_items = items[:limit]

        for item in target_items:
            title = item.find('title').text if item.find('title') is not None else "No Title"
            pub_date = item.find('pubDate').text if item.find('pubDate') is not None else ""
            news_list.append(f"- {title} ({pub_date})")
            
        return "\n".join(news_list)
    except Exception as e:
        print(f"News fetch error: {e}")
        return "뉴스 데이터를 가져오는 중 오류가 발생했습니다."

def generate_ai_report(company_id: int, report_id: int):
    """
    Generate VC investment report (Long Form - Modular Approach):
    Calls AI multiple times to ensure length and stability.
    """
    from ..db import SessionLocal
    
    db = SessionLocal()
    
    try:
        _append_report_log(report_id, "start")
        version_tag = "report_service_version=2026-01-23-1900"
        print(version_tag)
        _append_report_log(report_id, version_tag)
        # 1. Fetch Company Info
        company = db.execute(text("SELECT name_ko, corp_code, stock_code, sector_name FROM company WHERE company_id = :cid"), 
                             {"cid": company_id}).fetchone()
        
        if not company:
            db.execute(text("UPDATE report_request SET status = 'FAILED' WHERE report_id = :rid"),
                      {"rid": report_id})
            db.commit()
            return
        
        name = company[0]
        corp_code = company[1]
        ticker = company[2] or "N/A"
        sector = company[3] or "Technology"

        # Update status to RUNNING immediately
        db.execute(text("UPDATE report_request SET status = 'RUNNING' WHERE report_id = :rid"), {"rid": report_id})
        db.commit()

        print(f"Starting report generation for {name} (ID: {report_id})...")
        _append_report_log(report_id, "company_loaded")

        # Ensure financial marts are populated for this company (DART on-demand)
        print(f"Checking financials for {name}...")
        def _ensure_financials_async():
            try:
                with SessionLocal() as fin_db:
                    ensure_financials_for_company(fin_db, company_id, corp_code, ticker, name)
            except Exception as e:
                print(f"Warning: Financials check failed, proceeding with available data. Error: {e}")
                _append_report_log(report_id, f"financials_failed:{e}")

        try:
            executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
            executor.submit(_ensure_financials_async)
            executor.shutdown(wait=False, cancel_futures=False)
        except Exception as e:
            print(f"Warning: Financials async start failed: {e}")
            _append_report_log(report_id, f"financials_async_start_failed:{e}")
        
        print(f"Fetching news for {name}...")
        _append_report_log(report_id, "news_fetch_start")
        # 2. Fetch News (Limit to 15 for speed)
        news_summary = fetch_google_news(name, limit=15) 
        today_str = datetime.date.today().strftime("%Y-%m-%d")
        _append_report_log(report_id, "news_fetch_done")
        
        # 3. Setup AI
        api_key = settings.GOOGLE_API_KEY or os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            raise ValueError("GOOGLE_API_KEY not configured")
        
        genai.configure(api_key=api_key)
        # Use faster model for speed
        model_id = os.environ.get("GEMINI_MODEL_ID", "models/gemini-2.0-flash")
        model = genai.GenerativeModel(model_id)
        _append_report_log(report_id, f"model_ready:{model_id}")
        fin_summary_rows = db.execute(text("""
            SELECT fiscal_year, revenue, op_income, net_income, assets, equity
            FROM fs_mart_annual
            WHERE company_id = :cid
            ORDER BY fiscal_year DESC
            LIMIT 3
        """), {"cid": company_id}).fetchall()

        fin_ratio_rows = db.execute(text("""
            SELECT fiscal_year, op_margin, roe, debt_ratio
            FROM fs_ratio_mart
            WHERE company_id = :cid AND period_type = 'ANNUAL'
            ORDER BY fiscal_year DESC
            LIMIT 3
        """), {"cid": company_id}).fetchall()

        def _to_float(value):
            try:
                return float(value) if value is not None else None
            except (TypeError, ValueError):
                return None

        def _fmt_krw_100m(value):
            if value is None:
                return '-'
            return f"{value / 100000000:,.1f}억"

        def _fmt_pct(value):
            if value is None:
                return '-'
            return f"{value:.1f}%"

        summary_by_year = {}
        for row in fin_summary_rows:
            summary_by_year[row.fiscal_year] = {
                'revenue': _to_float(row.revenue),
                'op_income': _to_float(row.op_income),
                'net_income': _to_float(row.net_income),
                'assets': _to_float(row.assets),
                'equity': _to_float(row.equity),
            }

        ratio_by_year = {}
        for row in fin_ratio_rows:
            ratio_by_year[row.fiscal_year] = {
                'op_margin': _to_float(row.op_margin),
                'roe': _to_float(row.roe),
                'debt_ratio': _to_float(row.debt_ratio),
            }

        data_note = ''

        if corp_code:
            fs_rows = db.execute(text("""
                SELECT period_end, item_name, value, consolidated_flag
                FROM financial_statement
                WHERE corp_code = :cc
                ORDER BY period_end DESC, consolidated_flag DESC
            """), {"cc": corp_code}).fetchall()

            fs_by_year: dict[int, list] = {}
            for row in fs_rows:
                if not row.period_end:
                    continue
                year = row.period_end.year
                fs_by_year.setdefault(year, []).append(row)

            def _pick_value(items, keywords):
                for item in items:
                    name = str(item.item_name or '')
                    for key in keywords:
                        if key in name:
                            return _to_float(item.value)
                return None

            revenue_keys = ["매출", "영업수익", "수익(매출액)", "이자수익", "보험수익"]
            op_income_keys = ["영업이익", "영업이익(손실)"]
            net_income_keys = ["당기순이익", "당기순이익(손실)", "순이익"]
            assets_keys = ["자산총계", "총자산"]
            equity_keys = ["자본총계", "총자본", "자기자본"]

            def _pick_with_fallback(rows_all, rows_consolidated, keys):
                return _pick_value(rows_consolidated, keys) or _pick_value(rows_all, keys)

            for year in sorted(fs_by_year.keys(), reverse=True)[:3]:
                rows = fs_by_year[year]
                consolidated_rows = [r for r in rows if r.consolidated_flag]

                if year not in summary_by_year:
                    summary_by_year[year] = {}

                summary_by_year[year] = {
                    'revenue': summary_by_year[year].get('revenue') or _pick_with_fallback(rows, consolidated_rows, revenue_keys),
                    'op_income': summary_by_year[year].get('op_income') or _pick_with_fallback(rows, consolidated_rows, op_income_keys),
                    'net_income': summary_by_year[year].get('net_income') or _pick_with_fallback(rows, consolidated_rows, net_income_keys),
                    'assets': summary_by_year[year].get('assets') or _pick_with_fallback(rows, consolidated_rows, assets_keys),
                    'equity': summary_by_year[year].get('equity') or _pick_with_fallback(rows, consolidated_rows, equity_keys),
                }

            if fs_by_year:
                data_note = 'DART 재무 데이터 기준으로 작성되었습니다.'

        if summary_by_year:
            for year, summary in summary_by_year.items():
                revenue = summary.get('revenue')
                op_income = summary.get('op_income')
                net_income = summary.get('net_income')
                assets = summary.get('assets')
                equity = summary.get('equity')
                if year not in ratio_by_year:
                    ratio_by_year[year] = {}
                ratio_by_year[year] = {
                    'op_margin': ratio_by_year[year].get('op_margin') or ((op_income / revenue * 100) if revenue else None),
                    'roe': ratio_by_year[year].get('roe') or ((net_income / equity * 100) if equity else None),
                    'debt_ratio': ratio_by_year[year].get('debt_ratio') or (((assets - equity) / equity * 100) if assets is not None and equity else None),
                }

        try:
            years = _ensure_latest_years(db, corp_code)
        except ValueError as exc:
            print(f"Financial years fallback: {exc}")
            years = _candidate_years()[:3]
            if not data_note:
                data_note = '재무 데이터가 부족하여 일부 항목은 - 로 표시됩니다.'
        for y in years:
            summary_by_year.setdefault(y, {})
            ratio_by_year.setdefault(y, {})

        def _build_table(headers, rows):
            header = '| ' + ' | '.join(headers) + ' |'
            separator = '| ' + ' | '.join([':---'] + [':---:' for _ in headers[1:]]) + ' |'
            body = ['| ' + ' | '.join(row) + ' |' for row in rows]
            return '\n'.join([header, separator] + body)

        financial_rows = [
            ['매출액'] + [_fmt_krw_100m(summary_by_year.get(y, {}).get('revenue')) for y in years],
            ['영업이익'] + [_fmt_krw_100m(summary_by_year.get(y, {}).get('op_income')) for y in years],
            ['순이익'] + [_fmt_krw_100m(summary_by_year.get(y, {}).get('net_income')) for y in years],
            ['자산총계'] + [_fmt_krw_100m(summary_by_year.get(y, {}).get('assets')) for y in years],
            ['자본총계'] + [_fmt_krw_100m(summary_by_year.get(y, {}).get('equity')) for y in years],
        ]
        financial_table = _build_table(['구분'] + [f'{y}년' for y in years], financial_rows)

        ratio_rows = [
            ['영업이익률'] + [_fmt_pct(ratio_by_year.get(y, {}).get('op_margin')) for y in years],
            ['ROE'] + [_fmt_pct(ratio_by_year.get(y, {}).get('roe')) for y in years],
            ['부채비율'] + [_fmt_pct(ratio_by_year.get(y, {}).get('debt_ratio')) for y in years],
        ]
        ratio_table = _build_table(['지표'] + [f'{y}년' for y in years], ratio_rows)

        def _has_fin_values():
            for summary in summary_by_year.values():
                if any(v is not None for v in summary.values()):
                    return True
            for ratio in ratio_by_year.values():
                if any(v is not None for v in ratio.values()):
                    return True
            return False

        has_fin_data = _has_fin_values()
        if not has_fin_data:
            data_note = '재무 데이터가 부족하여 일부 항목은 - 로 표시됩니다.'

        dart_filings_md = "## 부록: 최근 공시 50건\n공시 데이터가 없습니다."
        if corp_code:
            dart_rows = db.execute(text("""
                SELECT filing_date, filing_type, title, rcp_no
                FROM dart_filing
                WHERE corp_code = :cc
                ORDER BY filing_date DESC
                LIMIT 50
            """), {"cc": corp_code}).fetchall()
            if dart_rows:
                lines = []
                for row in dart_rows:
                    filing_date = str(row.filing_date) if row.filing_date else "-"
                    filing_type = row.filing_type or "-"
                    title = row.title or "-"
                    rcp_no = row.rcp_no or "-"
                    lines.append(f"- {filing_date} [{filing_type}] {title} (rcp_no: {rcp_no})")
                dart_filings_md = "## 부록: 최근 공시 50건\n" + "\n".join(lines)
        

        # --- Prepare Prompts ---
        print("Preparing prompts for parallel generation...")
        _append_report_log(report_id, "prompts_ready")
        
        prompt1 = f"""
당신은 VC 수석 심사역입니다. {name}({ticker}) 보고서의 **요약 및 재무**를 작성하십시오.
정량/정성 분석을 포함한 전문적인 톤으로 작성하십시오.
본문에 "제1부/제2부/제3부" 같은 파트 표기는 절대 포함하지 마십시오.

[목차]
# 투자 검토 보고서: {name}

## 1. 투자 요약 (Executive Summary)
(A4 1페이지 분량으로 핵심 요약: 시장 기회, 리스크, 투자 포인트, 결론)

## 2. 재무 실적 및 지표 심층 분석

### 2.1 재무상태표 및 손익계산서 요약
{financial_table}

### 2.2 주요 재무 지표
{ratio_table}
{data_note}

### 2.3 재무 분석
(재무제표 수치 변화의 원인과 질적 변화를 5문단 이상으로 분석)

### 2.4 매출 및 이익 성장
(사업/제품별 성장 동력과 구조적 요인을 분석)

### 2.5 수익성/효율성 분석
(마진 변화, 비용 구조, 운영 효율 개선 가능성)

### 2.6 재무 건전성
(부채비율, 현금흐름, 유동성 관점에서 위험 요인 평가)

마크다운으로 작성하십시오.
"""

        prompt2 = f"""
당신은 VC 수석 심사역입니다. {name}({ticker}) 보고서의 **사업 모델**을 작성하십시오.
정량/정성 근거를 포함해 서술하십시오.
본문에 "제1부/제2부/제3부" 같은 파트 표기는 절대 포함하지 마십시오.

[목차]
## 3. 사업 모델 및 경쟁 우위

### 3.1 핵심 사업 구조와 수익원
(주요 제품/서비스, 매출 구성, 수익 구조를 설명)

### 3.2 시장 포지션 및 경쟁 환경
(시장 규모, 성장률, 경쟁사 대비 차별점)

### 3.3 성장 전략 및 제품 로드맵
(신규 시장/제품 확장 전략과 실행 가능성)

### 3.4 운영 역량 및 ESG
(생산/공급망/인재/ESG 관점에서 평가)

### 3.5 결론: 사업 모델의 지속 가능성
(향후 3~5년 관점에서 지속성 평가)

**뉴스 참고:**
{news_summary}

마크다운으로 작성하십시오.
"""

        prompt3 = f"""
당신은 VC 수석 심사역입니다. {name}({ticker}) 보고서의 **리스크/기회/결론**을 작성하십시오.
정확하고 구체적인 항목으로 작성하십시오.
본문에 "제1부/제2부/제3부" 같은 파트 표기는 절대 포함하지 마십시오.

[목차]
## 4. 리스크 및 기회 요인

### 4.1 리스크 요인 (Risk Factors)
(각 항목 3문장 이상, 정확히 8개)
- [리스크 1]
- [리스크 2]
- [리스크 3]
- [리스크 4]
- [리스크 5]
- [리스크 6]
- [리스크 7]
- [리스크 8]

### 4.2 기회 요인 (Opportunity Factors)
(각 항목 3문장 이상, 정확히 8개)
- [기회 1]
- [기회 2]
- [기회 3]
- [기회 4]
- [기회 5]
- [기회 6]
- [기회 7]
- [기회 8]

## 5. 최종 투자 결론 및 모니터링 포인트

### 5.1 최종 투자 결론
(투자 의견과 근거를 3문단 이상)

### 5.2 모니터링 포인트
(정확히 8개 항목)
- [포인트 1]
- [포인트 2]
- [포인트 3]
- [포인트 4]
- [포인트 5]
- [포인트 6]
- [포인트 7]
- [포인트 8]

마크다운으로 작성하십시오.
"""

        prompt4 = f"""
당신은 VC 수석 심사역입니다. {name}({ticker}) 보고서의 **뉴스 인사이트**를 작성하십시오.
최근 뉴스 기준으로 16개 항목을 정리하십시오.
본문에 "제1부/제2부/제3부" 같은 파트 표기는 절대 포함하지 마십시오.

**뉴스 참고:**
{news_summary}

[목차]
## 6. 최근 뉴스 및 인사이트
(각 항목에 뉴스 제목과 의미를 요약)
- [뉴스 1]
- [뉴스 2]
- [뉴스 3]
- [뉴스 4]
- [뉴스 5]
- [뉴스 6]
- [뉴스 7]
- [뉴스 8]
- [뉴스 9]
- [뉴스 10]
- [뉴스 11]
- [뉴스 12]
- [뉴스 13]
- [뉴스 14]
- [뉴스 15]
- [뉴스 16]

마크다운으로 작성하십시오.
"""

        # --- Parallel Generation ---
        print("Generating report parts in parallel...")
        _append_report_log(report_id, "generation_start")
        
        genai_timeout_sec = int(os.environ.get("GENAI_TIMEOUT_SEC", "45"))

        def _generate_part(prompt, timeout_sec: int | None = None):
            try:
                # Create a local model instance for thread safety if needed (though genai is stateless)
                local_model = genai.GenerativeModel(model_id)
                resp = local_model.generate_content(
                    prompt,
                    request_options={"timeout": timeout_sec or genai_timeout_sec},
                )
                return resp.text
            except Exception as e:
                print(f"Part generation failed: {e}")
                _append_report_log(report_id, f"part_failed:{type(e).__name__}:{e}")
                return ""

        part_timeout_sec = int(os.environ.get("GENAI_PART_TIMEOUT_SEC", "60"))
        executor = concurrent.futures.ThreadPoolExecutor(max_workers=4)
        try:
            future1 = executor.submit(_generate_part, prompt1)
            future2 = executor.submit(_generate_part, prompt2)
            future3 = executor.submit(_generate_part, prompt3)
            future4 = executor.submit(_generate_part, prompt4)

            futures = [future1, future2, future3, future4]
            raw_parts = ["", "", "", ""]
            for idx, fut in enumerate(futures):
                try:
                    raw_parts[idx] = fut.result(timeout=part_timeout_sec)
                except concurrent.futures.TimeoutError:
                    print(f"Part generation timed out after {part_timeout_sec}s (part {idx + 1})")
                    fut.cancel()
                except Exception as exc:
                    print(f"Part generation failed: {exc}")
                    raw_parts[idx] = ""

            raw_part1, raw_part2, raw_part3, raw_part4 = raw_parts
        finally:
            executor.shutdown(wait=False, cancel_futures=True)

        def _raw_len(text: str) -> int:
            if not text:
                return 0
            return len(re.sub(r"\\s+", "", text))

        raw_lengths = {
            "part1": _raw_len(raw_part1),
            "part2": _raw_len(raw_part2),
            "part3": _raw_len(raw_part3),
            "part4": _raw_len(raw_part4),
        }
        _append_report_log(report_id, f"raw_lengths:{raw_lengths}")

        retry_threshold = 200
        retry_targets = [k for k, v in raw_lengths.items() if v < retry_threshold]
        if retry_targets:
            retry_timeout_sec = int(os.environ.get("GENAI_RETRY_TIMEOUT_SEC", "90"))
            retry_map = {
                "part1": prompt1,
                "part2": prompt2,
                "part3": prompt3,
                "part4": prompt4,
            }
            for key in retry_targets:
                _append_report_log(report_id, f"retry_start:{key}")
                retried = _generate_part(retry_map[key], timeout_sec=retry_timeout_sec)
                _append_report_log(report_id, f"retry_done:{key}:{_raw_len(retried)}")
                if key == "part1":
                    raw_part1 = retried
                elif key == "part2":
                    raw_part2 = retried
                elif key == "part3":
                    raw_part3 = retried
                elif key == "part4":
                    raw_part4 = retried

        _append_report_log(report_id, "generation_done")
        # --- Process Part 1 ---
        part1 = clean_markdown(raw_part1)
        part1 = remove_part_labels(part1)
        part1 = normalize_headings(part1)
        part1 = normalize_report_header(part1, name, ticker)
        part1 = _inject_financial_tables(part1, financial_table, ratio_table, data_note)

        # --- Process Part 2 ---
        part2 = remove_part_labels(clean_markdown(raw_part2))
        part2 = normalize_headings(part2)
        part2 = normalize_part_header(part2, "제2부 (사업 모델)")
        part2 = ensure_heading(part2, "3. 사업 모델 및 경쟁 우위")

        # --- Process Part 3 ---
        part3 = remove_part_labels(clean_markdown(raw_part3))
        part3 = normalize_headings(part3)
        part3 = normalize_part_header(part3, "제3부 (리스크/기회/결론)")
        part3 = ensure_heading(part3, "4. 리스크 및 기회 요인")

        # --- Process Part 4 ---
        part4 = remove_part_labels(clean_markdown(raw_part4))
        part4 = normalize_headings(part4)
        part4 = normalize_part_header(part4, "제4부 (뉴스 인사이트)")
        part4 = ensure_heading(part4, "6. 최근 뉴스 및 인사이트")

        def _meaningful_len(text: str) -> int:
            if not text:
                return 0
            lines = []
            for line in text.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("#"):
                    continue
                if stripped.startswith("|"):
                    continue
                if re.match(r"^[:\\-\\s|]+$", stripped):
                    continue
                if re.match(r"^[-*]\\s*\\[.+\\]$", stripped):
                    continue
                lines.append(stripped)
            cleaned = re.sub(r"\\s+", "", " ".join(lines))
            return len(cleaned)

        part_lengths = {
            "part1": _meaningful_len(part1),
            "part2": _meaningful_len(part2),
            "part3": _meaningful_len(part3),
            "part4": _meaningful_len(part4),
        }
        _append_report_log(report_id, f"part_lengths:{part_lengths}")

        min_len = {
            "part1": 200,
            "part2": 200,
            "part3": 200,
            "part4": 120,
        }
        low_parts = [k for k, v in part_lengths.items() if v < min_len[k]]
        if low_parts:
            _append_report_log(report_id, f"low_content_parts:{low_parts}")
            raise ValueError(f"AI generation returned insufficient content for: {', '.join(low_parts)}")

        if not any([part1.strip(), part2.strip(), part3.strip(), part4.strip()]):
            raise ValueError("AI generation returned empty content for all parts.")

        # Combine all parts
        full_report = f"{part1}\n\n{part2}\n\n{part3}\n\n{part4}\n\n{dart_filings_md}"
        full_report = remove_part_labels(full_report)
        full_report = normalize_headings(full_report)
        full_report = _dedupe_financial_sections(full_report)
        full_report = _dedupe_section_heading(full_report, "## 3. 사업 모델 및 경쟁 우위")
        _validate_report_output(full_report)
        
        # 4. Save as Markdown
        root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../.."))
        artifacts_dir = os.path.join(root_dir, "artifacts", "reports")
        os.makedirs(artifacts_dir, exist_ok=True)
        
        md_path = os.path.join(artifacts_dir, f"report_{report_id}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(full_report)
        print(f"Markdown saved: {md_path}")
        _append_report_log(report_id, f"md_saved:{md_path}")
        
        # 5. Convert to DOCX
        docx_path = os.path.join(artifacts_dir, f"report_{report_id}.docx")
        markdown_to_docx_converter(full_report, docx_path, name, ticker)
        print(f"DOCX saved: {docx_path}")
        _append_report_log(report_id, f"docx_saved:{docx_path}")
        
        # 6. Update DB status to DONE
        db.execute(text("UPDATE report_request SET status = 'DONE' WHERE report_id = :rid"),
                  {"rid": report_id})
        db.commit()
        _append_report_log(report_id, "done")
        
        print(f"Report {report_id} generated successfully!")
        
    except Exception as e:
        print(f"Report generation failed: {e}")
        _append_report_log(report_id, f"failed:{e}")
        try:
            root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../.."))
            error_log_path = os.path.join(root_dir, "artifacts", "report_error.log")
            with open(error_log_path, "a", encoding="utf-8") as f:
                f.write(f"Timestamp: {datetime.datetime.now()}\n")
                f.write(f"Report ID: {report_id}, Company ID: {company_id}\n")
                f.write(f"Error: {str(e)}\n")
                import traceback
                f.write(traceback.format_exc())
                f.write("-" * 50 + "\n")
        except:
            pass
        db.rollback()
        db.execute(text("UPDATE report_request SET status = 'FAILED' WHERE report_id = :rid"),
                  {"rid": report_id})
        db.commit()
    
    finally:
        db.close()

def clean_markdown(text):
    """Clean up markdown code blocks"""
    text = text.strip()
    if text.startswith("```markdown"):
        text = text[11:]
    elif text.startswith("```"):
        text = text[3:]
    
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()

def _normalize_text(text: str) -> str:
    if not text:
        return text
    text = text.replace("\u200b", "").replace("\ufeff", "")
    return unicodedata.normalize("NFKC", text)

def remove_part_labels(markdown_text: str) -> str:
    """Remove '제1부/제2부/...' labels from content."""
    normalized = _normalize_text(markdown_text)
    cleaned = re.sub(r"(?m)^\\s*(제\\s*)?\\d+\\s*부\\s*[:\\-–]*\\s*", "", normalized)
    cleaned = re.sub(r"(제\\s*)?\\d+\\s*부", "", cleaned)
    return cleaned

def _inject_financial_tables(markdown_text: str, financial_table: str, ratio_table: str, data_note: str) -> str:
    """Force-inject 2.1/2.2 tables so they don't get rewritten or dropped by the model."""
    block = "\n".join([
        "### 2.1 재무상태표 및 손익계산서 요약",
        financial_table,
        "",
        "### 2.2 주요 재무 지표",
        ratio_table,
        data_note or "",
        "",
    ])

    # Remove any existing 2.1/2.2 blocks to avoid duplication.
    cleaned = re.sub(r"(?s)###\\s*2\\.1.*?(?=###\\s*2\\.\\d+|##\\s*3\\.|$)", "", markdown_text)
    cleaned = re.sub(r"(?s)###\\s*2\\.2.*?(?=###\\s*2\\.\\d+|##\\s*3\\.|$)", "", cleaned)

    # Otherwise insert right after the 2.x main header if present.
    heading_pattern = r"(?m)^##\\s*2\\..*$"
    if re.search(heading_pattern, cleaned):
        return re.sub(heading_pattern, lambda m: f"{m.group(0)}\n\n{block}", cleaned, count=1)

    return f"{cleaned}\n\n{block}"

def _safe_ratio(numerator: float | None, denominator: float | None) -> float | None:
    if numerator is None or denominator in (None, 0):
        return None
    return (numerator / denominator) * 100

def _pick_metric(existing: tuple[float | None, bool] | None, value: float | None, consolidated: bool) -> tuple[float | None, bool]:
    if value is None:
        return existing or (None, consolidated)
    if existing is None:
        return (value, consolidated)
    if existing[1] and not consolidated:
        return existing
    if not existing[1] and consolidated:
        return (value, consolidated)
    return (value, consolidated)

def build_marts_from_financial_statement(db, company_id: int, corp_code: str, years: list[int]):
    if not corp_code or not years:
        return
    period_ends = [datetime.date(y, 12, 31) for y in years]
    rows = db.execute(
        text("""
            SELECT period_end, item_name, value, consolidated_flag
            FROM financial_statement
            WHERE corp_code = :cc
              AND period_end = ANY(:pends)
        """),
        {"cc": corp_code, "pends": period_ends}
    ).fetchall()

    metric_map = {
        "revenue": ["매출액", "영업수익", "매출", "수익(매출액)", "이자수익", "보험수익"],
        "op_income": ["영업이익", "영업이익(손실)"],
        "net_income": ["당기순이익", "당기순이익(손실)", "연결당기순이익", "순이익"],
        "assets": ["자산총계", "자산총액", "총자산"],
        "liabilities": ["부채총계", "부채총액", "총부채"],
        "equity": ["자본총계", "자본총액", "총자본"],
    }

    by_year: dict[int, dict[str, tuple[float | None, bool]]] = {}
    for period_end, item_name, value, consolidated_flag in rows:
        year = period_end.year
        if year not in by_year:
            by_year[year] = {}
        for metric, names in metric_map.items():
            if item_name in names:
                prev = by_year[year].get(metric)
                by_year[year][metric] = _pick_metric(prev, float(value) if value is not None else None, bool(consolidated_flag))

    for year, metrics in by_year.items():
        revenue = metrics.get("revenue", (None, False))[0]
        op_income = metrics.get("op_income", (None, False))[0]
        net_income = metrics.get("net_income", (None, False))[0]
        assets = metrics.get("assets", (None, False))[0]
        liabilities = metrics.get("liabilities", (None, False))[0]
        equity = metrics.get("equity", (None, False))[0]

        db.execute(
            text("""
                INSERT INTO fs_mart_annual (company_id, fiscal_year, revenue, op_income, net_income, assets, liabilities, equity, generated_at)
                VALUES (:cid, :fy, :rev, :op, :net, :assets, :liab, :eq, NOW())
                ON CONFLICT (company_id, fiscal_year)
                DO UPDATE SET
                    revenue = EXCLUDED.revenue,
                    op_income = EXCLUDED.op_income,
                    net_income = EXCLUDED.net_income,
                    assets = EXCLUDED.assets,
                    liabilities = EXCLUDED.liabilities,
                    equity = EXCLUDED.equity,
                    generated_at = NOW()
            """),
            {
                "cid": company_id,
                "fy": year,
                "rev": revenue,
                "op": op_income,
                "net": net_income,
                "assets": assets,
                "liab": liabilities,
                "eq": equity,
            }
        )

        db.execute(
            text("""
                INSERT INTO fs_ratio_mart (company_id, period_type, fiscal_year, fiscal_quarter, op_margin, roe, debt_ratio, generated_at)
                VALUES (:cid, 'ANNUAL', :fy, 4, :op_margin, :roe, :debt_ratio, NOW())
                ON CONFLICT (company_id, period_type, fiscal_year, fiscal_quarter)
                DO UPDATE SET
                    op_margin = EXCLUDED.op_margin,
                    roe = EXCLUDED.roe,
                    debt_ratio = EXCLUDED.debt_ratio,
                    generated_at = NOW()
            """),
            {
                "cid": company_id,
                "fy": year,
                "op_margin": _safe_ratio(op_income, revenue),
                "roe": _safe_ratio(net_income, equity),
                "debt_ratio": _safe_ratio(liabilities, equity),
            }
        )

    db.commit()

def _resolve_corp_code(db, corp_code: str | None, stock_code: str | None, name_ko: str | None) -> str | None:
    if corp_code:
        return corp_code
    if stock_code:
        # Heuristic: preferred shares often end with 5 (e.g., 006405 -> 006400)
        if stock_code.endswith("5"):
            base_code = f"{stock_code[:-1]}0"
            row = db.execute(
                text("SELECT corp_code FROM company WHERE stock_code = :sc AND corp_code IS NOT NULL"),
                {"sc": base_code}
            ).fetchone()
            if row and row[0]:
                return row[0]
    if name_ko and name_ko.endswith("우"):
        base_name = name_ko[:-1]
        row = db.execute(
            text("SELECT corp_code FROM company WHERE name_ko = :name AND corp_code IS NOT NULL"),
            {"name": base_name}
        ).fetchone()
        if row and row[0]:
            return row[0]
    return None

def _latest_years_for_company(db, corp_code: str | None) -> list[int]:
    if not corp_code:
        return []
    rows = db.execute(
        text("""
            SELECT DISTINCT EXTRACT(YEAR FROM period_end)::int AS y
            FROM financial_statement
            WHERE corp_code = :cc AND period_end IS NOT NULL
            ORDER BY y DESC
            LIMIT 3
        """),
        {"cc": corp_code}
    ).fetchall()
    return [row[0] for row in rows]

def _candidate_years() -> list[int]:
    current_year = datetime.date.today().year
    # DART annuals are typically available up to (current_year - 2).
    base_year = current_year - 2
    # Look back 8 years to ensure we find at least 3 valid years of data even if recent ones are missing or delayed.
    return [base_year - i for i in range(8)]

def _ensure_latest_years(db, corp_code: str, required: int = 3) -> list[int]:
    years = _latest_years_for_company(db, corp_code)
    if len(years) < required:
        raise ValueError(f"Insufficient DART financial years: {len(years)} (need {required})")
    return years[:required]

def ensure_financials_for_company(db, company_id: int, corp_code: str | None, stock_code: str | None, name_ko: str | None):
    corp_code = _resolve_corp_code(db, corp_code, stock_code, name_ko)
    if not corp_code:
        return

    candidate_years = _candidate_years()

    def _existing_years(years: list[int]) -> list[int]:
        rows = db.execute(
            text("""
                SELECT DISTINCT EXTRACT(YEAR FROM period_end)::int
                FROM financial_statement
                WHERE corp_code = :cc AND period_end IS NOT NULL
                  AND EXTRACT(YEAR FROM period_end)::int = ANY(:years)
            """),
            {"cc": corp_code, "years": years}
        ).fetchall()
        return [row[0] for row in rows]

    def _missing_years(years: list[int]) -> list[int]:
        existing_set = set(_existing_years(years))
        return [y for y in years if y not in existing_set]

    def _fetch_from_dart(years: list[int]) -> None:
        if not years:
            return
        import sys
        import os as _os
        ingest_path = _os.path.abspath(_os.path.join(_os.path.dirname(__file__), "../../../../services/ingest"))
        if ingest_path not in sys.path:
            sys.path.append(ingest_path)
        from ingest.dart_financials_loader import fetch_and_save_company_financials
        fetch_and_save_company_financials(corp_codes=[corp_code], years=years)

    # Optimized DART fetch logic: Stop if we have 3 years of data
    
    # 1. Check existing years
    existing = _existing_years(candidate_years)
    
    # If we already have 3 years, we can stop here.
    if len(existing) >= 3:
        # Check if the latest expected year is missing?
        # For speed, we skip this if we have enough historical data.
        # But let's be slightly robust: if we have 2021,2022,2023 but not 2024, it's okay unless user demands 2024.
        print(f"Skipping DART fetch: Found {len(existing)} years in DB for {name_ko}.")
        latest_years = _ensure_latest_years(db, corp_code)
        build_marts_from_financial_statement(db, company_id, corp_code, latest_years)
        return

    # 2. Fetch only missing years, but stop once we reach 3 total years
    missing_candidates = sorted([y for y in candidate_years if y not in existing], reverse=True)
    
    # How many more do we need?
    needed = 3 - len(existing)
    to_fetch = missing_candidates[:needed + 2] # Fetch a few more just in case recent ones are empty
    
    if to_fetch:
        try:
            print(f"Fetching missing DART years for {name_ko}: {to_fetch}")
            _fetch_from_dart(to_fetch)
        except Exception as e:
            print(f"DART financials fetch failed: {e}") 

    latest_years = _ensure_latest_years(db, corp_code)
    build_marts_from_financial_statement(db, company_id, corp_code, latest_years)

def normalize_section_headings(markdown_text):
    """섹션 번호 헤더의 레벨을 일관되게 정규화"""
    lines = markdown_text.splitlines()
    normalized = []
    for line in lines:
        stripped = line.lstrip()
        if stripped.startswith("### "):
            title = stripped[4:].strip()
            if re.match(r"^\d+\.\s", title) and not re.match(r"^\d+\.\d+", title):
                line = line.replace("### ", "## ", 1)
        normalized.append(line)
    return "\n".join(normalized)

def normalize_headings(markdown_text: str) -> str:
    """Normalize heading markers and remove list bullets that leak into headings."""
    keywords = [
        "투자 요약",
        "재무 실적",
        "사업 모델",
        "경쟁 우위",
        "리스크",
        "기회",
        "최종 투자 결론",
        "모니터링 포인트",
        "최근 뉴스",
        "인사이트",
        "시장 포지션",
        "시장 환경",
        "재무상태표",
        "손익계산서",
    ]
    lines = markdown_text.splitlines()
    out = []
    for line in lines:
        raw = line.strip()
        if raw.endswith("**") and raw.startswith("#"):
            raw = raw[:-2].rstrip()
        if raw.startswith(("-", "*", "•")):
            candidate = raw.lstrip("-*•").strip()
            if re.match(r"^\d+(\.\d+)?\s", candidate) or any(k in candidate for k in keywords):
                raw = candidate

        m = re.match(r"^\*\*(.+)\*\*$", raw)
        if m:
            title = m.group(1).strip()
            if re.match(r"^\d+\.\d+\s", title):
                out.append(f"### {title}")
                continue
            if re.match(r"^\d+\.\s", title):
                out.append(f"## {title}")
                continue
            out.append(f"### {title}")
            continue

        if re.match(r"^\d+\.\d+\s", raw):
            out.append(f"### {raw}")
            continue
        if re.match(r"^\d+\.\s", raw):
            out.append(f"## {raw}")
            continue

        out.append(raw if raw else line)
    return "\n".join(out)

def _dedupe_financial_sections(markdown_text: str) -> str:
    """Remove duplicated 2.1/2.2 blocks that can appear after injection."""
    lines = markdown_text.splitlines()
    out = []
    seen_21 = 0
    skipping = False
    for line in lines:
        if line.startswith("### 2.1 재무상태표") or line.startswith("### 2.1 "):
            seen_21 += 1
            if seen_21 > 1:
                skipping = True
                continue
        if skipping:
            if line.startswith("## 3."):
                skipping = False
                out.append(line)
            continue
        out.append(line)
    return "\n".join(out)

def _dedupe_section_heading(markdown_text: str, heading: str) -> str:
    """Keep the first occurrence of a section heading and drop duplicates."""
    lines = markdown_text.splitlines()
    out = []
    seen = False
    for line in lines:
        if line.startswith(heading):
            if not seen:
                out.append(heading)
                seen = True
            continue
        out.append(line)
    return "\n".join(out)

def ensure_heading(markdown_text: str, heading: str) -> str:
    """Ensure a top-level heading exists in the block."""
    if re.search(rf"(?m)^##\\s*{re.escape(heading)}\\b", markdown_text):
        return markdown_text
    return f"## {heading}\n\n{markdown_text}".strip()

def normalize_report_header(markdown_text, name, ticker):
    """첫 줄 헤더를 고정 포맷으로 정규화"""
    markdown_text = normalize_headings(normalize_section_headings(_normalize_text(markdown_text)))
    lines = markdown_text.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines:
        first = lines[0].lstrip()
        if first.startswith("#"):
            lines.pop(0)
        elif "투자 검토 보고서" in first:
            lines.pop(0)
    while lines and not lines[0].strip():
        lines.pop(0)
    ticker_display = ticker if ticker else "-"
    fixed_header = f"# 투자 검토 보고서: {name} ({ticker_display})"
    return "\n".join([fixed_header, ""] + lines)

def normalize_part_header(markdown_text, _part_title):
    """파트 제목 제거: 본문 헤더만 유지"""
    markdown_text = normalize_section_headings(markdown_text)
    lines = markdown_text.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines and lines[0].lstrip().startswith("#"):
        lines.pop(0)
    while lines and not lines[0].strip():
        lines.pop(0)
    return "\n".join(lines)

def markdown_to_docx_converter(markdown_text, output_path, company_name, ticker):
    """留덊겕?ㅼ슫??DOCX濡?蹂??(H4 諛??쒕툕 ?뱀뀡 吏??媛뺥솕)"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    import re
    
    doc = Document()
    
    # ?섏씠吏 ?ㅼ젙
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = markdown_text.split('\n')
    in_table = False
    table_rows = []
    
    for line in lines:
        line_stripped = line.strip()
        
        # H1 (Title)
        if line_stripped.startswith('# ') and not line_stripped.startswith('## '):
            text = line_stripped[2:].strip()
            p = doc.add_paragraph(text)
            p.style = 'Heading 1'
            for run in p.runs:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(14, 165, 233) # Blue
                
        # H2 (Main Section)
        elif line_stripped.startswith('## '):
            text = line_stripped[3:].strip()
            p = doc.add_paragraph(text)
            p.style = 'Heading 2'
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0) # Black
                
        # H3 (Sub Section 2.1, 2.2 etc)
        elif line_stripped.startswith('### '):
            text = line_stripped[4:].strip()
            p = doc.add_paragraph(text)
            p.style = 'Heading 3'
            for run in p.runs:
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(60, 60, 60) # Dark Gray
                
        # H4 (Deep Detail if needed)
        elif line_stripped.startswith('#### '):
            text = line_stripped[5:].strip()
            p = doc.add_paragraph(text)
            p.style = 'Heading 4'
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.bold = True
        
        # Table
        elif line_stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = [line_stripped]
            else:
                table_rows.append(line_stripped)
        
        # Table End check
        elif not line_stripped.startswith('|') and in_table:
            _create_table(doc, table_rows)
            in_table = False
            table_rows = []
            
            # Process current line after table
            _process_line(doc, line_stripped)
            
        else:
            _process_line(doc, line_stripped)
    
    if in_table and table_rows:
        _create_table(doc, table_rows)
    
    doc.save(output_path)

def _validate_report_output(markdown_text: str) -> None:
    """Validate report output for mandatory rules."""
    normalized = _normalize_text(markdown_text)
    if re.search(r"제\s*\d+\s*부", normalized):
        raise ValueError("Report output contains forbidden part labels (제N부).")

    table_ok = False
    for line in normalized.splitlines():
        if line.strip().startswith("| 구분 |") or line.strip().startswith("| 지표 |"):
            # Count columns between pipes
            cols = [c for c in line.split("|") if c.strip()]
            # Expect at least 4 columns: label + 3 years
            if len(cols) >= 4:
                table_ok = True
            else:
                raise ValueError("Financial table has fewer than 3 years.")
    if not table_ok:
        raise ValueError("Financial tables not found in report output.")

def _process_line(doc, line):
    """Helper to process normal lines, bullets, etc."""
    import re
    if not line:
        return
        
    # Bullet point
    if line.startswith('o ') or line.startswith('- '):
        text = line[2:].strip()
        p = doc.add_paragraph(text, style='List Bullet')
        
    # Numbered list
    elif re.match(r'^\d+[\.\)]\s', line):
        text = re.sub(r'^\d+[\.\)]\s', '', line)
        p = doc.add_paragraph(text, style='List Number')
        
    # Bold text processing
    elif '**' in line:
        p = doc.add_paragraph()
        parts = line.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True
    else:
        doc.add_paragraph(line)

def _create_table(doc, table_rows):
    """Create DOCX table from markdown table rows"""
    if len(table_rows) < 2:
        return
    
    header = [c.strip() for c in table_rows[0].split('|')[1:-1]]
    data_rows = []
    
    # Skip header(0) and separator(1)
    for row in table_rows[2:]:
        if row.strip():
            cells = [c.strip() for c in row.split('|')[1:-1]]
            data_rows.append(cells)
    
    if not data_rows:
        return
    
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.autofit = False

    # Keep column widths consistent across tables (align 2.1 and 2.2)
    try:
        from docx.shared import Inches
        total_inches = 6.27  # A4 width 8.27 - 1in margins on both sides
        col_count = len(header)
        if col_count >= 4:
            first_inches = 1.8
        else:
            first_inches = total_inches / max(col_count, 1)
        remaining_inches = total_inches - first_inches
        other_inches = remaining_inches / max(col_count - 1, 1)
        widths = [first_inches] + [other_inches] * max(col_count - 1, 0)
        for idx, w in enumerate(widths):
            width = Inches(w)
            for cell in table.columns[idx].cells:
                cell.width = width
    except Exception:
        pass
    
    # Header
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < len(table.rows[r_idx + 1].cells):
                table.rows[r_idx + 1].cells[c_idx].text = cell_text

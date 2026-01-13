"""
기존 DOCX 템플릿에 플레이스홀더를 추가하는 스크립트
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 원본 템플릿 로드
template_path = r"c:\ProjectCode\stockmanager\docs\vc투자메모.docx"
doc = Document(template_path)

# 플레이스홀더 매핑 (기존 텍스트 → 플레이스홀더)
replacements = {
    "삼성전자": "{{COMPANY_NAME}}",
    "005930": "{{TICKER}}",
    "2026년 01월 10일": "{{DATE}}",
}

# 모든 문단에서 교체
for para in doc.paragraphs:
    for old, new in replacements.items():
        if old in para.text:
            # 스타일 유지하면서 텍스트만 교체
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

# 표에서도 교체
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in replacements.items():
                    if old in para.text:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)

# 템플릿으로 저장
output_path = r"c:\ProjectCode\stockmanager\docs\vc투자메모_template.docx"
doc.save(output_path)
print(f"Template created: {output_path}")
print("\nPlaceholders added:")
for old, new in replacements.items():
    print(f"  {old} → {new}")
